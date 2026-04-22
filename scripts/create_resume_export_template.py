from pathlib import Path

from docx import Document


SOURCE = Path("content/templates/xiaoming-resume-template.docx")
TARGET = Path("content/templates/xiaoming-resume-export-template.docx")


PARAGRAPH_MAP = {
    0: "{{profile_name}}",
    1: "{{profile_headline}}",
    2: "{{profile_contact}}",
    3: "{{profile_location}}",
    4: "{{profile_links}}",
    6: "{{education_1_header}}",
    7: "{{education_1_detail}}",
    8: "{{education_2_header}}",
    9: "{{education_2_detail}}",
    11: "{{fulltime_1_header}}",
    12: "{{fulltime_1_mix}}",
    13: "{{fulltime_1_bullet_1}}",
    14: "{{fulltime_1_bullet_2}}",
    15: "{{fulltime_1_bullet_3}}",
    16: "{{fulltime_1_bullet_4}}",
    17: "{{fulltime_1_bullet_5}}",
    18: "{{fulltime_1_bullet_6}}",
    19: "{{fulltime_2_header}}",
    20: "{{fulltime_2_mix}}",
    21: "{{fulltime_2_bullet_1}}",
    22: "{{fulltime_2_bullet_2}}",
    23: "{{fulltime_2_bullet_3}}",
    24: "{{fulltime_2_bullet_4}}",
    25: "{{fulltime_2_bullet_5}}",
    26: "{{fulltime_2_bullet_6}}",
    28: "{{intern_1_header}}",
    29: "{{intern_1_bullet_1}}",
    30: "{{intern_1_bullet_2}}",
    31: "{{intern_1_bullet_3}}",
    32: "{{intern_2_header}}",
    33: "{{intern_2_bullet_1}}",
    34: "{{intern_2_bullet_2}}",
    35: "{{intern_2_bullet_3}}",
    36: "{{intern_2_bullet_4}}",
    41: "{{venture_header}}",
    42: "{{venture_bullet_1}}",
    43: "{{venture_bullet_2}}",
    44: "{{venture_bullet_3}}",
    45: "{{venture_bullet_4}}",
    47: "{{skill_1}}",
    48: "{{skill_2}}",
    49: "{{skill_3}}",
    51: "{{strength_1}}",
    52: "{{strength_2}}",
    53: "{{strength_3}}",
    54: "{{profile_manual}}",
    55: "{{profile_learning}}",
    85: "{{en_strength_1}}",
    86: "{{en_strength_2}}",
    87: "{{en_strength_3}}",
    88: "{{en_strength_4}}",
    90: "{{en_fulltime_1_header}}",
    91: "{{en_fulltime_1_bullet_1}}",
    92: "{{en_fulltime_1_bullet_2}}",
    93: "{{en_fulltime_1_bullet_3}}",
    94: "{{en_fulltime_2_header}}",
    95: "{{en_fulltime_2_bullet_1}}",
    96: "{{en_fulltime_2_bullet_2}}",
    97: "{{en_fulltime_2_bullet_3}}",
    99: "{{en_intern_1_header}}",
    100: "{{en_intern_1_bullet_1}}",
    101: "{{en_intern_1_bullet_2}}",
    102: "{{en_intern_2_header}}",
    103: "{{en_intern_2_bullet_1}}",
    104: "{{en_intern_2_bullet_2}}",
    105: "{{en_intern_3_header}}",
    106: "{{en_intern_3_bullet_1}}",
    107: "{{en_intern_3_bullet_2}}",
    109: "{{en_venture_header}}",
    110: "{{en_venture_bullet_1}}",
    111: "{{en_venture_bullet_2}}",
    112: "{{en_venture_bullet_3}}",
    114: "{{en_skill_1}}",
    115: "{{en_skill_2}}",
    117: "{{en_education_1_header}}",
    118: "{{en_education_1_detail}}",
    119: "{{en_education_2_header}}",
}


def main():
    doc = Document(str(SOURCE))
    for index, placeholder in PARAGRAPH_MAP.items():
        if index < len(doc.paragraphs):
            doc.paragraphs[index].text = placeholder
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TARGET))
    print(f"created {TARGET}")


if __name__ == "__main__":
    main()
